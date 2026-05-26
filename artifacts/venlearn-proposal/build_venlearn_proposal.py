from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "artifacts" / "venlearn-proposal"
OUT_DOCX = OUT_DIR / "Venlearn_School_Management_Software_Proposal.docx"
LOGO_PATH = Path("/Users/ucheumeh/Downloads/Venlearn Logo/Venlearn O&B.png")
COVER_PATH = Path("/Users/ucheumeh/Downloads/venlearn proposal cover page 2.png")

BLUE = RGBColor(38, 97, 172)
DARK_BLUE = RGBColor(22, 74, 134)
INK = RGBColor(28, 36, 52)
MUTED = RGBColor(77, 88, 104)
ORANGE = RGBColor(255, 128, 0)
BLACK = RGBColor(0, 0, 0)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "2661AC"
STRIPE_FILL = "F3F7FC"
BLUE_FILL = "EAF2FB"
ORANGE_FILL = "FFF3E6"
WHITE = RGBColor(255, 255, 255)
FONT_NAME = "Arial"
BODY_FONT_SIZE = 12


FEATURES = [
    (
        "Students Management",
        "Complete student biodata, class placement, admission status, documents, academic history, guardians, transfers, withdrawals, alumni records, promotions, emergency contacts, notes, tags, houses, clubs, and reports.",
    ),
    (
        "Staff Management",
        "Teacher and non-academic staff profiles, departments, subject and class assignments, employment records, contact details, attendance, workload, activity history, permission groups, directories, and exports.",
    ),
    (
        "Guardians Management",
        "Parent and guardian profiles linked to one or many students, relationship types, contact channels, pickup permissions, emergency contacts, portal access, communication history, household records, and family reporting.",
    ),
    (
        "Fee Collection",
        "Fee schedules, invoices, receipts, balances, discounts, payment plans, arrears, payment recording for cash, bank transfer, POS and online payments, reminders, debtor lists, and finance summaries.",
    ),
    (
        "Results & Report Cards",
        "Assessment scores, exam marks, grading scales, report card templates, online result checking, report card printing, score imports from CBT, Excel and other sources, class averages, teacher and principal remarks, approval workflows, class positions, downloadable reports, portal publishing, psychomotor notes, attendance summaries, and result analytics.",
    ),
    (
        "CBT: Offline & Online",
        "Question banks for objective, theory, fill-in-the-blank, image and audio questions, Microsoft Word imports, offline LAN exams, online tests, timing, randomized questions, shuffled options, exam windows, monitoring, switch alerts, automatic marking, manual review, result publishing, and analysis reports.",
    ),
    (
        "Digital Learning & eLibrary",
        "Notes, books, videos, assignments, past questions, study resources, class and subject organization, portal access, visibility controls, and engagement signals for popular or recently uploaded materials.",
    ),
    (
        "Attendance",
        "Daily, class, subject, hostel and staff attendance registers, late/absent/excused/custom statuses, parent notifications, term summaries, date-range reports, and exception patterns for leadership review.",
    ),
    (
        "Timetable",
        "Class, teacher, room, subject and period scheduling, conflict checks, student and staff views, substitutions, announcements, and reusable termly schedule structures.",
    ),
    (
        "Library Management",
        "Book cataloging, authors, categories, shelves, copies, ISBNs, availability, borrowing, renewals, reservations, overdue returns, fines, borrower histories, reminders, and inventory reports.",
    ),
    (
        "Hostel Management",
        "Hostel, dormitory, room, bed and resident allocation, boarding status, capacity tracking, hostel attendance, leave requests, incidents, welfare notes, emergency records, occupancy, availability and movement reports.",
    ),
    (
        "Event Calendar",
        "Academic calendars for terms, holidays, exams, open days, meetings, activities, recurring events, category visibility, reminders, and shared web and portal views.",
    ),
    (
        "Medicals and Incidents Reporting",
        "Medical profiles, allergies, conditions, medications, clinic notes, incident records, guardian notifications, follow-up actions, staff responsibility tracking, confidential access controls, and wellbeing reports.",
    ),
    (
        "Parents & Students Portal",
        "Secure family and learner access to results, report cards, invoices, receipts, attendance, assignments, learning resources, calendars, timetables, teacher updates, messages, and role-aware records.",
    ),
    (
        "Lesson Planner",
        "Lesson objectives, topics, activities, materials, assessments, subject/class/teacher/week/term organization, department review, approvals, reusable templates, resource attachments, and progress tracking.",
    ),
    (
        "Messaging and Communication",
        "Announcements, reminders, targeted notices for classes, departments, hostels, clubs, individuals or the whole school, message history, emergency alerts, delivery records, and reusable templates.",
    ),
    (
        "Inventory and Facility Management",
        "Supplies, equipment, assets, classrooms, stock levels, issue logs, returns, purchases, reorder visibility, maintenance requests, repairs, facility usage, accountability reports, and operational planning records.",
    ),
    (
        "Extracurricular Activities",
        "Clubs, societies, sports, houses, activity groups, participant lists, coordinators, schedules, events, attendance, achievements, awards, competitions, communication, and engagement reports.",
    ),
]


def set_run_font(run, name=FONT_NAME, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=190, bottom=110, end=190):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (
        ("top", top),
        ("start", start),
        ("left", start),
        ("bottom", bottom),
        ("end", end),
        ("right", end),
    ):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def style_cell_text(cell, bold=False, color=INK, size=BODY_FONT_SIZE):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.03
        for run in paragraph.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_para(doc, text="", *, style=None, bold=False, color=INK, size=BODY_FONT_SIZE, align=None, after=8, before=0, line=1.28):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r, size=BODY_FONT_SIZE, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r, size=BODY_FONT_SIZE, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    if level == 1:
        size, color, before, after = 17, BLUE, 26, 10
    elif level == 2:
        size, color, before, after = 14, BLUE, 18, 6
    else:
        size, color, before, after = 12.5, DARK_BLUE, 12, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def add_callout(doc, title, body, fill=BLUE_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=320, start=320, bottom=320, end=320)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=BODY_FONT_SIZE, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=BODY_FONT_SIZE, color=INK)
    add_para(doc, "", after=8)
    return table


def add_table(doc, headers, rows, widths, header_fill=HEADER_FILL, font_size=BODY_FONT_SIZE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    set_table_borders(table, color="C8D6E8", size="4")
    mark_header_row(table.rows[0])
    keep_row_together(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_fill)
        cell.text = h
        style_cell_text(cell, bold=True, color=WHITE, size=font_size)
    for row_index, row in enumerate(rows):
        added_row = table.add_row()
        keep_row_together(added_row)
        cells = added_row.cells
        for i, value in enumerate(row):
            tc_pr = cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(i, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cells[i], STRIPE_FILL)
            cells[i].text = value
            style_cell_text(cells[i], bold=(i == 0), size=font_size)
    add_para(doc, "", after=8)
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("Phone: +2348164177129  |  Email: info@veracone.com  |  Website: venlearn.com")
    set_run_font(r, size=9.5, color=BLUE, bold=True)
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("17, Ifeanyi Chinonye Crescent, Startimes Estate, Okota. Lagos.")
    set_run_font(r2, size=9, color=MUTED)


def add_letterhead(section):
    header = section.header
    for paragraph in header.paragraphs:
        paragraph.text = ""

    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [3600, 5760])
    remove_table_borders(table)
    mark_header_row(table.rows[0])
    logo_cell, contact_cell = table.rows[0].cells
    logo_p = logo_cell.paragraphs[0]
    logo_p.paragraph_format.space_after = Pt(0)
    logo_run = logo_p.add_run()
    if LOGO_PATH.exists():
        logo_shape = logo_run.add_picture(str(LOGO_PATH), width=Inches(1.75))
        logo_shape._inline.docPr.set("descr", "Venlearn logo in orange and blue.")
        logo_shape._inline.docPr.set("title", "Venlearn logo")
    else:
        set_run_font(logo_run, size=15, color=BLUE, bold=True)
        logo_run.text = "Venlearn"

    contact_p = contact_cell.paragraphs[0]
    contact_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_p.paragraph_format.space_after = Pt(0)
    contact_p.paragraph_format.line_spacing = 1.0
    contact_lines = [
        "Phone: +2348164177129",
        "Email: info@veracone.com  |  Website: venlearn.com",
        "17, Ifeanyi Chinonye Crescent, Startimes Estate, Okota. Lagos.",
    ]
    for idx, line in enumerate(contact_lines):
        if idx:
            contact_p.add_run().add_break()
        run = contact_p.add_run(line)
        set_run_font(run, size=8.4, color=BLUE if idx < 2 else MUTED, bold=idx < 2)

    bar = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(bar, [2800, 6560])
    remove_table_borders(bar)
    mark_header_row(bar.rows[0])
    set_cell_shading(bar.cell(0, 0), "FF8000")
    set_cell_shading(bar.cell(0, 1), "2661AC")
    for cell in bar.rows[0].cells:
        cell.text = ""
        set_cell_margins(cell, top=12, bottom=12, start=0, end=0)


def configure_cover_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)


def configure_body_section(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.95)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0.35)
    add_footer(section)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(BODY_FONT_SIZE)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title = styles["Title"]
    title.font.name = FONT_NAME
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = BLUE

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.bold = True


def configure_doc(doc):
    configure_body_section(doc.sections[0])
    configure_styles(doc)


def compress_image(src, dst, width_px=1300):
    img = Image.open(src).convert("RGB")
    if img.width > width_px:
        ratio = width_px / img.width
        img = img.resize((width_px, int(img.height * ratio)))
    img.save(dst, "JPEG", quality=85, optimize=True)


def add_cover_rule(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2600, 6760])
    remove_table_borders(table)
    mark_header_row(table.rows[0])
    set_cell_shading(table.cell(0, 0), "FF8000")
    set_cell_shading(table.cell(0, 1), "2661AC")
    for cell in table.rows[0].cells:
        cell.text = ""
        set_cell_margins(cell, top=20, start=0, bottom=20, end=0)
    add_para(doc, "", after=18)


def add_section_page_break(doc):
    doc.add_page_break()


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_cover_section(doc.sections[0])
    configure_styles(doc)

    if COVER_PATH.exists():
        cover_paragraph = doc.add_paragraph()
        cover_paragraph.paragraph_format.space_before = Pt(0)
        cover_paragraph.paragraph_format.space_after = Pt(0)
        cover_paragraph.paragraph_format.line_spacing = 1.0
        cover_run = cover_paragraph.add_run()
        cover_shape = cover_run.add_picture(str(COVER_PATH))
        cover_shape.width = Inches(8.5)
        cover_shape.height = Inches(11)
        cover_shape._inline.docPr.set("descr", "Venlearn proposal cover page with school building, Venlearn branding, and Veracone contact details.")
        cover_shape._inline.docPr.set("title", "Venlearn proposal cover page")
    else:
        add_para(doc, "Venlearn Complete School Management Software", bold=True, color=BLUE, size=28, align=WD_ALIGN_PARAGRAPH.CENTER, before=120, after=8, line=1.08)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_body_section(body_section)

    add_heading(doc, "Executive Recommendation", 1)
    add_callout(
        doc,
        "Recommendation",
        "Adopt the full Venlearn Complete feature set as the school's central operating system so administration, academics, finance, CBT, parent engagement, welfare, hostel, inventory, communication, and campus operations are connected from the outset.",
    )
    add_para(
        doc,
        "Venlearn brings the everyday work of a school into one connected workspace. Instead of keeping student data in registers, fees in separate ledgers, report cards in spreadsheets, announcements in messaging apps, and exam records in isolated tools, the school can maintain one reliable source of truth across administrators, teachers, bursars, parents, students, and leadership.",
    )
    add_para(
        doc,
        "For school administrators, the value is not only that Venlearn has many features. The value is that the features are connected: a student record links to guardians, fees, attendance, results, timetable, hostel, incidents, learning resources, and communication history. This makes decisions faster, reduces avoidable errors, improves parent confidence, and gives leadership clearer visibility into what is happening each term.",
    )

    add_heading(doc, "Why Schools Need Venlearn", 1)
    add_bullet(doc, "Administrators spend too much time reconciling paper registers, Excel sheets, fee books, lesson notes, printed report cards, and scattered files.")
    add_bullet(doc, "Parents expect timely updates on attendance, fees, events, results, assignments, and school notices.")
    add_bullet(doc, "School leaders need reliable reports on student population, fee collection, staff activity, academic performance, attendance, assets, safety, and operations.")
    add_bullet(doc, "Teachers need simple workflows for attendance, lesson planning, score entry, learning resources, and class communication.")
    add_bullet(doc, "Schools with unstable internet still need software that can work offline for core operations and CBT.")
    add_bullet(doc, "Schools whose current software providers are inadequate need a stronger upgrade with broader features, better usability, reliable support, and room for future growth.")
    add_bullet(doc, "Schools that want to adopt computer-based testing but have not figured out the right process need a ready CBT system that supports offline LAN exams, online tests, question imports, monitoring, marking, and result publishing.")
    add_para(
        doc,
        "Venlearn directly addresses these issues by combining cloud, offline/local deployment, role-based access, online and offline CBT, parent and student portals, fee management, report cards, communication, and operational modules in one platform.",
    )

    add_heading(doc, "The Venlearn Advantage", 1)
    add_table(
        doc,
        ["Advantage", "What it means for the school"],
        [
            ("One source of truth", "Student, staff, guardian, finance, academic, attendance and operations records live in one workspace, reducing duplication, errors and reporting delays."),
            ("Role-aware access", "Admins, staff, parents and students see only the records and actions meant for them, so sensitive data stays protected while each user gets a simpler experience."),
            ("Offline and cloud options", "Venlearn can run online, offline over a local area network, or with sync where required, supporting schools with inconsistent internet or local-control needs."),
            ("CBT built for schools", "Question banks, Word imports, offline LAN tests, online tests, monitoring, timing, marking and result publishing help schools move beyond basic forms."),
            ("Better user interface and smooth user experience", "A modern, role-aware workspace makes daily tasks easier for administrators, staff, parents and students, improving adoption and reducing frustration."),
            ("Parent confidence", "Parents can access fees, receipts, attendance, results, announcements, assignments, calendars and messages, reducing repetitive calls to the office."),
            ("Operational breadth", "Hostel, medicals, incidents, inventory, facilities, events, library and activities are covered, so the system supports campus life beyond exams and fees."),
            ("Available customer support", "Guidance is available for setup, migration, training, troubleshooting, workflow questions and day-to-day account support."),
        ],
        [2700, 6660],
    )

    add_heading(doc, "Complete Feature Inventory", 1)
    add_para(
        doc,
        "The following modules are included in the recommended full Venlearn adoption. Implementing all core features gives the school one connected operating system instead of separate tools for records, fees, academics, exams, communication, welfare, and operations.",
    )
    features_by_name = {name: description for name, description in FEATURES}
    feature_groups = [
        (
            "Feature Inventory: Records, Access and Daily Administration",
            ["Students Management", "Staff Management", "Guardians Management", "Parents & Students Portal", "Fee Collection", "Attendance"],
        ),
        (
            "Feature Inventory: Academics, Assessment and Learning",
            ["Results & Report Cards", "CBT: Offline & Online", "Digital Learning & eLibrary", "Lesson Planner", "Timetable", "Library Management"],
        ),
        (
            "Feature Inventory: Communication, Welfare and Campus Operations",
            ["Messaging and Communication", "Event Calendar", "Medicals and Incidents Reporting", "Hostel Management", "Inventory and Facility Management", "Extracurricular Activities"],
        ),
    ]
    for group_index, (group_title, names) in enumerate(feature_groups):
        if group_index:
            add_section_page_break(doc)
        add_heading(doc, group_title, 2)
        feature_rows = [(name, features_by_name[name]) for name in names]
        add_table(doc, ["Module", "What the module covers"], feature_rows, [2300, 7060])

    add_section_page_break(doc)
    add_heading(doc, "Benefits by Stakeholder", 1)
    add_heading(doc, "For School Owners, Directors and Principals", 2)
    add_bullet(doc, "Clearer visibility into enrolment, attendance, fee collection, academic performance, staff work, student welfare, and operational activities.")
    add_bullet(doc, "Better accountability because workflows, approvals, communication, records, and reports are tied to users and dates.")
    add_bullet(doc, "A more modern school image during admissions, parent engagement, assessments, and reporting.")
    add_heading(doc, "For Administrators and Bursars", 2)
    add_bullet(doc, "Faster record lookup, fee balance checks, receipt handling, arrears follow-up, timetable updates, and report preparation.")
    add_bullet(doc, "Reduced dependence on paper files and manual spreadsheet reconciliation.")
    add_heading(doc, "For Teachers", 2)
    add_bullet(doc, "Simpler access to class lists, attendance, lesson plans, score entry, learning resources, timetables, messages, and assessments.")
    add_bullet(doc, "Less repetitive formatting work when results and report cards are generated through a structured workflow.")
    add_heading(doc, "For Parents and Guardians", 2)
    add_bullet(doc, "Secure visibility into their wards' fees, receipts, attendance, results, events, assignments, learning materials, and notices.")
    add_bullet(doc, "More timely reminders and fewer missed updates about fees, events, attendance concerns, incidents, or school activities.")
    add_heading(doc, "For Students", 2)
    add_bullet(doc, "Access to timetables, assignments, eLibrary materials, CBT activities, results, announcements, and learning resources.")

    add_section_page_break(doc)
    add_heading(doc, "Cost Advantage and Income Improvement", 1)
    add_para(
        doc,
        "Venlearn is not only an administrative expense. Used properly, it helps the school reduce avoidable costs, protect revenue, improve collections, and create a more attractive experience for parents who are deciding where to enrol their children.",
    )
    add_table(
        doc,
        ["Opportunity", "How Venlearn Creates Financial Value"],
        [
            ("Transparent fee payment", "Invoices, receipts, balances, arrears and reminders make fee payment easier to track, reduce confusion, and help the school detect leakages or unrecorded payments."),
            ("Timely collections", "Automated fee reminders and parent portal visibility encourage earlier payment, reduce follow-up workload, and minimize losses from forgotten or disputed balances."),
            ("Time saved by teachers", "CBT marking, score imports, report card generation and structured result workflows reduce time spent marking scripts, entering scores and formatting reports, turning staff hours into direct value for the school."),
            ("Higher parent trust", "Parents feel relief when they can see their wards' attendance, results, assignments, fees, announcements and school activities, which improves satisfaction and supports retention."),
            ("Improved enrolment", "Online school operations, portals, enquiries, admissions workflows and remote access make it easier for families anywhere in the world to begin or complete enrolment conversations."),
        ],
        [2500, 6860],
    )
    add_section_page_break(doc)
    add_heading(doc, "Cost Advantage and Income Improvement (Continued)", 1)
    add_table(
        doc,
        ["Opportunity", "How Venlearn Creates Financial Value"],
        [
            ("Lower operational overhead", "One dashboard for records, fees, academics, hostel, inventory, welfare and communication reduces the need to hire separate dedicated heads or clerks to manually manage disconnected operations."),
            ("Reduced printing and paper costs", "Digital records, online results, downloadable report cards, messages, eLibrary materials and portal access reduce repeated printing, photocopying, physical file storage and manual distribution."),
            ("Better asset accountability", "Inventory and facility records help the school track equipment, supplies, repairs, stock levels and usage, reducing waste and unexplained losses."),
            ("Faster leadership decisions", "Dashboards and reports give school leaders a clearer view of enrolment, finance, attendance, academic performance and operations before problems become expensive."),
        ],
        [2500, 6860],
    )

    add_heading(doc, "Recommended Packaging and Pricing", 1)
    add_table(
        doc,
        ["Plan", "Included modules", "Price"],
        [
            (
                "Essentials",
                "Students, staff, guardians, attendance, timetable, event calendar, messaging and parent/student portal.",
                "NGN 2,000 per student per term.",
            ),
            (
                "Growth",
                "Everything in Essentials plus fee collection, results/report cards, lesson planner, eLibrary, library, medicals and incidents.",
                "NGN 3,000 per student per term.",
            ),
            (
                "Complete",
                "Everything in Growth plus CBT offline/online, hostel, inventory/facility management, extracurricular activities, offline/local deployment option and workflow setup support.",
                "NGN 4,000 per student per term.",
            ),
        ],
        [1700, 5360, 2300],
    )
    add_para(
        doc,
        "Recommended option: Complete. This gives the school the full Venlearn operating system from the start, including administration, finance, academics, CBT, portals, hostel, inventory, welfare, activities, offline/local deployment options and deeper rollout support.",
        bold=True,
        color=DARK_BLUE,
    )

    add_heading(doc, "Implementation and Onboarding Plan", 1)
    add_table(
        doc,
        ["Step", "What Happens"],
        [
            ("1. Walkthroughs & Demo", "Venlearn demonstrates the full system, answers administrator questions, confirms the school's priorities, and shows the workflows for leadership, bursary, teachers, parents and students."),
            ("2. Setup and Onboarding", "The school structure, sessions, terms, classes, subjects, departments, fee items, roles, permissions, portals and selected templates are configured."),
            ("3. Data Migration", "Venlearn helps import existing student, staff, guardian, class, fee and academic records from clean templates or supported exports from the school's current system."),
            ("4. Training and Support", "Administrators, bursars, teachers, leadership and support staff receive role-based training, with support available for launch questions and ongoing usage."),
        ],
        [2300, 7060],
    )

    add_heading(doc, "Training Plan", 1)
    add_bullet(doc, "Leadership orientation: dashboard visibility, reports, approvals, term review, data protection and adoption expectations.")
    add_bullet(doc, "Administrator training: student records, staff records, guardians, classes, admissions, promotions, timetable, messaging and general settings.")
    add_bullet(doc, "Bursary training: fee schedules, invoices, receipts, balances, discounts, arrears, payment recording, reminders and finance reports.")
    add_bullet(doc, "Teacher training: attendance, lesson planning, score entry, report remarks, assignments, resources, timetable, class communication and CBT participation.")
    add_bullet(doc, "Parent-support training: helping parents access portals, understand receipts, results, attendance, assignments, notices and common account questions.")
    add_bullet(doc, "Technical/admin handover: roles, permissions, backups or sync approach where applicable, escalation process and support contacts.")

    add_heading(doc, "Data Protection, Access and Reliability", 1)
    add_para(
        doc,
        "Venlearn should be configured around the principle that each person sees only what they need. Administrators can control access for staff, parents and students using role-based permissions. Sensitive medical, incident, finance and academic records should be restricted to approved users. For schools that require offline use, Venlearn can be deployed on a local area network so operations and CBT can continue without constant internet access, with sync options where needed.",
    )
    add_bullet(doc, "Use role-based access for administrators, bursars, teachers, parents, students and support staff.")
    add_bullet(doc, "Maintain audit-friendly records for key workflows such as fees, report cards, attendance, communication and incidents.")
    add_bullet(doc, "Confirm backup, sync and device policies during setup, especially for offline/local deployments.")

    add_section_page_break(doc)
    add_heading(doc, "Frequently Asked Administrator Questions", 1)
    faq_rows = [
        ("Will staff find it difficult to use?", "The software is easy to use. However, Venlearn offers training and support programs to customers to help them easily navigate the platform."),
        ("Do we need all features at once?", "The recommendation is to procure the Complete feature set so every major school workflow is covered. Day-to-day activation can still be scheduled carefully during onboarding so staff are trained properly."),
        ("What happens if our internet is unreliable?", "Venlearn supports offline/local deployment for schools that need local area network access, especially for CBT and daily operations."),
        ("Can parents see only their own children?", "Yes. Parent and guardian access is linked to their wards and controlled through secure role-aware portal views."),
        ("Can staff access only what they are assigned to access?", "Yes. Venlearn uses role-based access controls so staff can be limited to assigned classes, subjects, departments, reports, approvals or operational responsibilities."),
        ("Can it generate report cards?", "Yes. Venlearn supports scores, grading scales, remarks, approvals, class positions, report templates, downloadable reports and portal-ready publishing."),
        ("Does it help collect fees?", "Yes. It supports fee schedules, invoices, balances, receipts, arrears, discounts, payment plans, reminders and finance summaries."),
        ("Can it replace Google Forms for exams?", "For many school exam scenarios, yes. Venlearn CBT supports offline LAN exams, online tests, question banks, timing, monitoring, marking and result publishing."),
        ("I have an existing software. How can I move my data?", "Venlearn helps with data migration. Existing student, staff, guardian, class, finance and academic records can be prepared, cleaned and imported from supported exports or approved templates."),
        ("If I start with a smaller package, can I upgrade later?", "Yes. A school can start with Essentials or Growth and upgrade later to add more modules, users, deployment options and support as needs grow."),
        ("Do you offer discount for long yearly commitment?", "Yes. Venlearn offers discounts for yearly plans, with discounts as high as 25 percent depending on the package and commitment."),
        ("What if I am not satisfied?", "You can cancel your subscription anytime and still be able to download your school data as Excel files."),
        ("Do you support schools with different branches?", "Yes. Venlearn supports multi-branch or multi-campus schools, with centralized leadership visibility and branch-aware records, users, reports and operations."),
        ("Do you offer building or overhauling a school website?", "Yes. Venlearn can support building a new school website or overhauling an existing one. Website work comes at a fixed one-time cost separate from the software subscription."),
        ("Can we request custom changes?", "Yes. Custom reports, forms or workflow changes can be reviewed and quoted based on feasibility, timeline and product fit."),
        ("How long will onboarding take?", "Onboarding can be completed in one week, provided the school shares the required records, confirms responsible staff, and attends the scheduled training sessions."),
        ("What support is available after launch?", "Support should include troubleshooting, workflow guidance, refresher training, launch stabilization and periodic review with school administrators."),
    ]
    faq_groups = [
        ("Ease of Use, Access and Core Workflows", faq_rows[:5]),
        ("Results, Migration, Pricing and Commitments", faq_rows[5:11]),
        ("Support, Branches and Custom Work", faq_rows[11:]),
    ]
    for group_index, (group_title, rows) in enumerate(faq_groups):
        if group_index:
            add_section_page_break(doc)
            add_heading(doc, "Frequently Asked Administrator Questions (Continued)", 1)
        add_heading(doc, group_title, 2)
        add_table(doc, ["Question", "Answer"], rows, [3000, 6360])

    add_section_page_break(doc)
    add_heading(doc, "Success Measures for the First Term", 1)
    add_bullet(doc, "At least 90 percent of active student, staff and guardian records are complete and searchable.")
    add_bullet(doc, "Attendance, fee balances and selected academic workflows are used consistently by assigned staff.")
    add_bullet(doc, "Report cards are generated through Venlearn for the selected classes or school-wide rollout.")
    add_bullet(doc, "Parents receive agreed notices, fee reminders, attendance updates or portal access without relying only on manual calls.")
    add_bullet(doc, "Leadership receives usable reports for enrolment, attendance, fees, academics and selected operations.")
    add_bullet(doc, "A post-launch review identifies remaining training needs, configuration changes and next modules to activate.")

    add_heading(doc, "Next Steps for Approval", 1)
    add_number(doc, "Approve a 60-minute demo for school leadership, administrators, bursars and selected teachers.")
    add_number(doc, "Approve the recommended Complete plan so all major Venlearn features are included in the school rollout.")
    add_number(doc, "Share current school size, number of campuses, internet conditions, desired modules, fee structure and existing record formats.")
    add_number(doc, "Agree the final quote, implementation timeline, data migration scope and training dates.")
    add_number(doc, "Nominate an internal project owner and module champions for administration, bursary, academics, ICT and parent support.")
    add_number(doc, "Begin setup, data collection and staff onboarding ahead of the target launch date.")
    add_para(doc, "", after=14)
    add_callout(
        doc,
        "Conversion close",
        "Venlearn gives the school a complete path to modern administration: one operating system for records, fees, academics, communication, assessment, welfare, hostel, inventory, activities, and campus life from the outset.",
    )

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_document()
    print(OUT_DOCX)
